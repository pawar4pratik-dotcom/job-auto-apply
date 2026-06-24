"""
resume_parser.py
─────────────────
Parses a resume (PDF or DOCX) into:
  1. raw_text   — full plain text, used as LLM context
  2. sections    — text split by detected headers (Experience, Education, Skills, ...)
  3. facts        — best-effort structured extraction (emails, phone, links, skill tokens)
"""

import re
import json
import os
from dataclasses import dataclass, field, asdict
from typing import Optional


def _extract_pdf_text(path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx_text(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf_text(path)
    elif ext in (".docx", ".doc"):
        return _extract_docx_text(path)
    else:
        raise ValueError(f"Unsupported resume format: {ext}")


SECTION_HEADERS = {
    "summary": [r"summary", r"profile", r"objective", r"about me"],
    "experience": [r"experience", r"employment history", r"work history"],
    "education": [r"education", r"academic background"],
    "skills": [r"skills", r"technical skills", r"core competencies", r"technologies"],
    "projects": [r"projects", r"personal projects"],
    "certifications": [r"certifications?", r"licenses?"],
    "languages": [r"languages"],
}


def split_sections(raw_text: str) -> dict:
    lines = raw_text.split("\n")
    sections = {}
    current_key = "header"
    buffer = []

    def flush():
        if buffer:
            sections.setdefault(current_key, "")
            sections[current_key] += "\n".join(buffer).strip() + "\n"

    for line in lines:
        stripped = line.strip()
        matched_key = None
        if stripped and len(stripped) < 40:
            lowered = stripped.lower().strip(":").strip()
            for key, patterns in SECTION_HEADERS.items():
                if any(re.fullmatch(p, lowered) or re.match(rf"^{p}$", lowered) for p in patterns):
                    matched_key = key
                    break
        if matched_key:
            flush()
            buffer = []
            current_key = matched_key
        else:
            buffer.append(line)
    flush()
    return sections


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}")
LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", re.I)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9\-_/]+", re.I)
YEARS_EXP_RE = re.compile(r"(\d+)\+?\s*(?:years|yrs)\s*(?:of)?\s*experience", re.I)


def extract_facts(raw_text: str, sections: dict) -> dict:
    facts = {}

    email_match = EMAIL_RE.search(raw_text)
    if email_match:
        facts["email"] = email_match.group(0)

    linkedin_match = LINKEDIN_RE.search(raw_text)
    if linkedin_match:
        facts["linkedin_url"] = linkedin_match.group(0)

    github_match = GITHUB_RE.search(raw_text)
    if github_match:
        facts["github_url"] = github_match.group(0)

    phone_match = PHONE_RE.search(raw_text)
    if phone_match:
        candidate = phone_match.group(0)
        if sum(c.isdigit() for c in candidate) >= 7:
            facts["phone"] = candidate.strip()

    years_match = YEARS_EXP_RE.search(raw_text)
    if years_match:
        facts["years_of_experience"] = years_match.group(1)

    skills_text = sections.get("skills", "")
    if skills_text:
        tokens = re.split(r"[,•|/\n]+", skills_text)
        facts["skills"] = [t.strip() for t in tokens if t.strip() and len(t.strip()) < 40]

    facts["summary_text"] = sections.get("summary", "").strip()
    return facts


@dataclass
class ParsedResume:
    raw_text: str = ""
    sections: dict = field(default_factory=dict)
    facts: dict = field(default_factory=dict)

    def to_dict(self):
        return asdict(self)


def parse_resume(path: str, cache_path: Optional[str] = None) -> ParsedResume:
    raw_text = extract_text(path)
    sections = split_sections(raw_text)
    facts = extract_facts(raw_text, sections)
    parsed = ParsedResume(raw_text=raw_text, sections=sections, facts=facts)

    if cache_path:
        os.makedirs(os.path.dirname(cache_path), exist_ok=True)
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(parsed.to_dict(), f, indent=2, ensure_ascii=False)

    return parsed


def load_cached_resume(cache_path: str) -> Optional[ParsedResume]:
    if not os.path.exists(cache_path):
        return None
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return ParsedResume(**data)
    except Exception:
        return None
